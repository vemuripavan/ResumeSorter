from .resumereader import ResumeReader
from docx import Document
import xml.dom.minidom

class DocxReader(ResumeReader):
    #Method which will parse the docx file and return text
    def readResume(self,file):
        document = Document(file)
        content = []
        
        headers = [x.blob.decode() for x in document.part.package.parts if x.partname.find('header')>0]
        for header in headers:        
            content = self.getChildText(xml.dom.minidom.parseString(header))
            
        footers = [x.blob.decode() for x in document.part.package.parts if x.partname.find('footer')>0]
        for footer in footers:        
            content = content + self.getChildText(xml.dom.minidom.parseString(footer))
        
        content = content + [para.text for para in document.paragraphs]
        # Read data from tables also
        tables = document.tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    content = content + [para.text for para in cell.paragraphs]
        
        
        return "\n".join(content)
    
    def getChildText(self, node):
        items = []
        if node.hasChildNodes:
            for cn in node.childNodes:
                items = items + self.getChildText(cn)
        if node.nodeValue  is not None: 
            items.append(node.nodeValue)
        return items
    
# Code for Unit Testing
"""
from docx import Document
import xml.dom.minidom

file="D:\deep\SPAN\Shikhsa\AI\ML\kaggle\Data Science_Final project\Resumes\Sharepoint\Adusumalli V.Ramanamma_3_iGrid_hyd_SP_N.docx"

document = Document(file)
content = []

headers = [x.blob.decode() for x in document.part.package.parts if x.partname.find('header')>0]
for header in headers:        
    content = self.getChildText(xml.dom.minidom.parseString(header))
    
footers = [x.blob.decode() for x in document.part.package.parts if x.partname.find('footer')>0]
for footer in footers:        
    content = content + self.getChildText(xml.dom.minidom.parseString(footer))


content = [para.text for para in document.paragraphs]
table = document.tables
content = content + [para.text for table in document.tables]
tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            content = content + [para.text for para in cell.paragraphs]


for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            print paragraph.text
            
tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)

text= "\n".join(content)
"""
